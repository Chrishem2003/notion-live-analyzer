from io import BytesIO
def docx_bytes(p):
 from docx import Document
 d=Document();d.add_heading(p.get("name","Your Name"),0);d.add_paragraph(p.get("headline",""));d.add_heading("Summary",1);d.add_paragraph(p.get("summary",""));d.add_heading("Skills",1);d.add_paragraph(", ".join(p.get("skills",[])));d.add_heading("Experience",1)
 for e in p.get("experience",[]):d.add_heading(f"{e.get('title','')} — {e.get('company','')}",2);d.add_paragraph(e.get("dates",""));d.add_paragraph(e.get("description",""))
 d.add_heading("Education",1);d.add_paragraph(p.get("education",""));b=BytesIO();d.save(b);return b.getvalue()
def pdf_bytes(p):
 from reportlab.platypus import SimpleDocTemplate,Paragraph
 from reportlab.lib.styles import getSampleStyleSheet
 from reportlab.lib.pagesizes import A4
 b=BytesIO();s=getSampleStyleSheet();story=[Paragraph(str(p.get("name","Your Name")),s["Title"]),Paragraph(str(p.get("headline","")),s["Heading2"]),Paragraph("Summary",s["Heading2"]),Paragraph(str(p.get("summary","")),s["BodyText"]),Paragraph("Skills",s["Heading2"]),Paragraph(", ".join(p.get("skills",[])),s["BodyText"]),Paragraph("Experience",s["Heading2"])]
 for e in p.get("experience",[]):story += [Paragraph(f"{e.get('title','')} — {e.get('company','')}",s["Heading3"]),Paragraph(str(e.get("description","")).replace("\n","<br/>"),s["BodyText"])]
 SimpleDocTemplate(b,pagesize=A4).build(story);return b.getvalue()
